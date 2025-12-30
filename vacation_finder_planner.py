import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '.')))
import streamlit as st
import openai
import requests
from datetime import datetime
from fpdf import FPDF
import io
import re
import pathlib
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown2
from bs4 import BeautifulSoup
from PIL import Image

# Set page config with custom icon
icon_path = os.path.join(os.path.dirname(__file__), 'DigitaL_Planner_App.png')
if os.path.exists(icon_path):
    try:
        icon_img = Image.open(icon_path)
        st.set_page_config(
            page_title="Vacation Finder & Planner",
            page_icon=icon_img,
            layout="wide"
        )
    except Exception as e:
        # Fallback to emoji if image fails to load
        st.set_page_config(
            page_title="Vacation Finder & Planner",
            page_icon="🗺️",
            layout="wide"
        )
else:
    # Fallback to emoji if image doesn't exist
    st.set_page_config(
        page_title="Vacation Finder & Planner",
        page_icon="🗺️",
        layout="wide"
    )

# Set your API keys
# Import API keys from config file (for local development)
# Falls back to environment variables (for GitHub/Streamlit Cloud deployment)
try:
    from config.config import OPENAI_API_KEY, SERP_API_KEY
except (ImportError, ModuleNotFoundError):
    # Fallback to environment variables for GitHub/Streamlit Cloud
    OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
    SERP_API_KEY = os.getenv('SERP_API_KEY')
    
    if not OPENAI_API_KEY or not SERP_API_KEY:
        st.error("❌ API keys not found. Please set OPENAI_API_KEY and SERP_API_KEY in config file or environment variables.")
        st.stop()

openai.api_key = OPENAI_API_KEY

SERPAPI_SEARCH_URL = "https://serpapi.com/search.json"

# Add this after imports, before any Streamlit UI code
st.markdown(
    '''
    <style>
    /* Remove the default Streamlit tab indicator bar and replace with a white underline */
    div[data-baseweb="tab-list"] > div {
        border-bottom: 2px solid #fff !important;
        border-radius: 0 !important;
        background: none !important;
        box-shadow: none !important;
    }
    /* Make the active tab label white and bold, with a white underline */
    div[data-baseweb="tab"] > button[aria-selected="true"] {
        color: #fff !important;
        font-weight: bold !important;
        border-bottom: 3px solid #fff !important;
        background: none !important;
        border-radius: 0 !important;
    }
    /* Make inactive tab labels gray */
    div[data-baseweb="tab"] > button[aria-selected="false"] {
        color: #bbb !important;
        background: none !important;
    }
    </style>
    ''',
    unsafe_allow_html=True
)

def search_travel_deals(start, dest, start_date, days, preferences):
    query = f"best travel deals {start} to {dest or 'anywhere'} {start_date} {days} days {preferences}"
    params = {
        "q": query,
        "api_key": SERP_API_KEY,
        "num": 5,
        "engine": "google",
        "hl": "en"
    }
    resp = requests.get(SERPAPI_SEARCH_URL, params=params)
    results = resp.json().get("organic_results", [])
    return results

def generate_itinerary(dest, start_date, days, preferences, deals, restaurant_preferences=""):
    context = "\n".join([f"{d['title']}: {d.get('snippet', '')} ({d.get('link', '')})" for d in deals])
    
    # Set restaurant recommendation text based on user input
    if restaurant_preferences and restaurant_preferences.strip():
        restaurant_text = f"Recommend good {restaurant_preferences.strip()} restaurants with good reviews along the way."
    else:
        restaurant_text = "By default, recommend good Indian, Thai, or Mexican restaurants with good reviews along the way."
    
    prompt = (
        f"Plan a detailed {days}-day vacation in {dest} starting on {start_date}. "
        f"Include daily activities, must-see places, and where to eat. "
        f"Optimize the route so that driving distance is minimized and the trip is convenient. "
        f"Choose a logical order for visiting places, making the route efficient but still interesting. "
        f"{restaurant_text} "
        f"For each night, recommend hotels with a price range around $200/night, with very good reviews and free breakfast, and provide links to book them if possible. "
        f"For each restaurant, provide a link to book or view the menu if possible. "
        f"Give a detailed, clear itinerary with places to see, what to do, and explanations for each. "
        f"Consider these preferences: {preferences}. "
        f"Use the following deals and links if relevant:\n{context}\n\n"
        f"Format as a day-by-day itinerary with links for booking."
    )
    response = openai.OpenAI(api_key=OPENAI_API_KEY).chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1500,
        temperature=0.7
    )
    return response.choices[0].message.content.strip()

def clean_and_fit_line(pdf, line, cell_width):
    # Only printable ASCII chars that fit in the cell
    safe_chars = []
    for char in line:
        if 32 <= ord(char) <= 126:  # printable ASCII
            if pdf.get_string_width(char) <= cell_width:
                safe_chars.append(char)
            # else: skip char that can't fit
    return ''.join(safe_chars)

def export_pdf(itinerary, deals):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    effective_width = pdf.w - 2 * pdf.l_margin
    pdf.multi_cell(0, 10, "Vacation Plan", align="C")
    pdf.ln(5)
    for line in itinerary.split('\n'):
        safe_line = clean_and_fit_line(pdf, line, effective_width)
        if safe_line.strip():
            pdf.multi_cell(0, 10, safe_line)
    pdf.ln(5)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 10, "Top Deals & Booking Links:")
    for d in deals:
        deal_line = f"- {d['title']}: {d.get('link', '')}"
        safe_deal_line = clean_and_fit_line(pdf, deal_line, effective_width)
        if safe_deal_line.strip():
            pdf.multi_cell(0, 10, safe_deal_line)
    pdf_buffer = io.BytesIO()
    pdf.output(pdf_buffer)
    pdf_buffer.seek(0)
    return pdf_buffer

def markdown_to_docx(md_text, doc=None):
    if doc is None:
        doc = Document()
    html = markdown2.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    for elem in soup.find_all(['h1', 'h2', 'h3', 'ul', 'ol', 'li', 'p', 'strong', 'b', 'a']):
        if elem.name == 'h1':
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == 'h2':
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == 'h3':
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == 'ul' or elem.name == 'ol':
            for li in elem.find_all('li'):
                doc.add_paragraph('• ' + li.get_text(), style='List Bullet')
        elif elem.name == 'p':
            doc.add_paragraph(elem.get_text())
        elif elem.name in ['strong', 'b']:
            para = doc.add_paragraph()
            run = para.add_run(elem.get_text())
            run.bold = True
        elif elem.name == 'a':
            para = doc.add_paragraph()
            run = para.add_run(elem.get_text() + ': ')
            # Add hyperlink using XML hack
            part = doc.part
            link = elem.get('href', '')
            if link:
                r_id = part.relate_to(link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '0000FF')
                rPr.append(color)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = link
                new_run.append(t)
                hyperlink.append(new_run)
                para._p.append(hyperlink)
    return doc

def export_docx(itinerary_md, deals):
    doc = Document()
    doc.add_heading('Vacation Plan', level=0)
    doc.add_heading('Itinerary', level=1)
    doc = markdown_to_docx(itinerary_md, doc)
    doc.add_heading('Top Deals & Booking Links', level=1)
    for d in deals:
        title = d.get('title', '')
        link = d.get('link', '')
        para = doc.add_paragraph(style='List Bullet')
        if link and link.startswith('http'):
            run = para.add_run(title + ': ')
            part = doc.part
            r_id = part.relate_to(link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = link
            new_run.append(t)
            hyperlink.append(new_run)
            para._p.append(hyperlink)
        else:
            para.add_run(f"{title}: {link}")
    for p in doc.paragraphs:
        p.paragraph_format.space_after = Pt(6)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Tabbed layout ---
tabs = st.tabs(["Vacation Finder & Planner", "Vacation Assistant"])

with tabs[0]:
    st.markdown(
        "<div style='background-color:#f8f9fa; padding: 0px 18px 18px 18px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.04); margin-top:-18px;'>",
        unsafe_allow_html=True
    )
    st.subheader("Vacation Finder & Planner", divider="rainbow")
    with st.form("vacation_form"):
        start = st.text_input("Start location (city or airport):", value=st.session_state.get('start', ''))
        dest = st.text_input("Destination (leave blank for 'anywhere'):", value=st.session_state.get('dest', ''))
        start_date = st.date_input("Start date:", min_value=datetime.today(), value=st.session_state.get('start_date', datetime.today()))
        days = st.number_input("Number of days:", min_value=1, max_value=30, value=st.session_state.get('days', 7))
        preferences = st.text_input("Preferences (e.g., cruise, city, nature, food, etc.):", value=st.session_state.get('preferences', ''))
        restaurant_prefs = st.text_input("Restaurant Preferences (e.g., Italian, French, Asian, etc.):", 
                                         value=st.session_state.get('restaurant_prefs', ''),
                                         help="By default, the app will recommend good Indian, Thai, or Mexican restaurants. Enter your preferred cuisine types here to customize.")
        submitted = st.form_submit_button("Find & Plan Vacation")

    if submitted:
        st.session_state['start'] = start
        st.session_state['dest'] = dest
        st.session_state['start_date'] = start_date
        st.session_state['days'] = days
        st.session_state['preferences'] = preferences
        st.session_state['restaurant_prefs'] = restaurant_prefs
        with st.spinner("Searching for deals and planning your trip..."):
            deals = search_travel_deals(start, dest, start_date, days, preferences)
            itinerary = generate_itinerary(dest or 'a great destination', start_date, days, preferences, deals, restaurant_prefs)
            st.session_state['vacation_itinerary'] = itinerary
            st.session_state['vacation_deals'] = deals

    # Display itinerary if it exists in session_state (persists across reruns)
    if 'vacation_itinerary' in st.session_state and 'vacation_deals' in st.session_state:
        st.subheader("Your Vacation Plan:")
        st.markdown(st.session_state['vacation_itinerary'])
        st.subheader("Top Deals & Booking Links:")
        for d in st.session_state['vacation_deals']:
            st.markdown(f"- [{d['title']}]({d.get('link', '')})")
        
        # Add separator line before Buy Me a Coffee
        st.markdown("---")
        
        # Buy Me a Coffee section - only shown after report is generated
        st.markdown(
            """
            <div style='display: flex; flex-direction: column; align-items: flex-end; margin-top: 20px; margin-bottom: 20px;'>
                <a href="https://www.buymeacoffee.com/balrajvishnu" target="_blank">
                    <img src="https://cdn.buymeacoffee.com/buttons/bmc-new-btn-logo.svg" alt="Buy Me a Coffee" style="height: 44px; width: 44px; border-radius: 50%; background: #FFDD00; padding: 4px;">
                </a>
                <a href="https://www.buymeacoffee.com/balrajvishnu" target="_blank" style="margin-top: 4px; color: #fff; background: rgba(0,0,0,0.7); padding: 2px 10px; border-radius: 8px; font-size: 1.08em; font-weight: 500; text-align: right; text-decoration: none;">
                    If you found this useful, buy me a coffee and support my work
                </a>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        # Show export button
        docx_file = export_docx(st.session_state['vacation_itinerary'], st.session_state['vacation_deals'])
        st.download_button(
            "Download Vacation Plan as DOCX", 
            docx_file, 
            file_name="vacation_plan.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    # Add spacing at the bottom for visual distinction between content and footer
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

with tabs[1]:
    st.markdown(
        "<div style='background-color:#f0f4fa; padding: 0px 18px 18px 18px; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.06); border-left: 3px solid #e0e0e0; margin-top:-18px;'>",
        unsafe_allow_html=True
    )
    st.header("Vacation Assistant 🗺️", divider="blue")
    st.markdown("<span style='color:#555'>Ask about places, distances, recommendations, etc.</span>", unsafe_allow_html=True)

    # Initialize chat history in session state
    if 'vacation_chat_history' not in st.session_state:
        st.session_state['vacation_chat_history'] = []

    # Display chat history
    for msg in st.session_state['vacation_chat_history']:
        if msg['role'] == 'user':
            st.markdown(f"<div style='margin-bottom:6px;'><b>You:</b> {msg['content']}</div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div style='margin-bottom:12px; color:#fff;'><b>Assistant:</b> {msg['content']}</div>", unsafe_allow_html=True)

    # Chat input
    with st.form("vacation_chat_form", clear_on_submit=True):
        user_question = st.text_area("Type your question:", key="vacation_chat_input", height=80, max_chars=500)
        chat_submitted = st.form_submit_button("Send")

    if chat_submitted and user_question.strip():
        # Add user message to history
        st.session_state['vacation_chat_history'].append({'role': 'user', 'content': user_question})
        # Compose messages for OpenAI
        system_prompt = (
            "You are a helpful vacation assistant. Answer questions about places to see, distances between cities, recommendations for restaurants, attractions, and travel tips. "
            "Be concise, friendly, and provide links or names if possible. If asked about distance, estimate in miles/km and travel time. If asked for recommendations, suggest well-reviewed options. "
            "If the user asks about a city or place, assume they are traveling or planning a trip."
        )
        messages = [
            {"role": "system", "content": system_prompt}
        ]
        for msg in st.session_state['vacation_chat_history']:
            messages.append({"role": msg['role'], "content": msg['content']})
        # Call OpenAI
        with st.spinner("Assistant is typing..."):
            try:
                response = openai.OpenAI(api_key=OPENAI_API_KEY).chat.completions.create(
                    model="gpt-4o",
                    messages=messages,
                    max_tokens=400,
                    temperature=0.5
                )
                answer = response.choices[0].message.content.strip()
            except Exception as e:
                answer = f"Sorry, there was an error: {e}"
        # Add assistant message to history
        st.session_state['vacation_chat_history'].append({'role': 'assistant', 'content': answer})
        st.experimental_rerun()
    st.markdown("</div>", unsafe_allow_html=True)

# Add clear separator line before footer
st.markdown("---")

# --- Footer with Privacy Policy and Disclaimer ---
# Initialize session state for privacy and disclaimer
if 'show_privacy' not in st.session_state:
    st.session_state['show_privacy'] = False
if 'show_disclaimer' not in st.session_state:
    st.session_state['show_disclaimer'] = False

# Footer button styling - horizontal, elegant, readable
footer_css = """
<style>
/* Style footer buttons - smaller font and gray color */
div[data-testid="column"]:has(button[key="vacation_privacy_btn"]) button,
div[data-testid="column"]:has(button[key="vacation_disclaimer_btn"]) button {
    font-size: 0.9em !important;
}
button:contains("Privacy Policy"), button:contains("Disclaimer") {
    color: #666 !important;
    font-size: 0.9em !important;
}
.footer-copyright {
    text-align: center;
    color: #888;
    font-size: 0.95em;
    margin-top: 1em;
    margin-bottom: 1em;
}
.privacy-main-title, .disclaimer-main-title {
    font-size: 1.4em;
    font-weight: bold;
    margin-bottom: 0.5em;
    color: #666;
}
.privacy-expander-content h2, .disclaimer-expander-content h2 {
    font-size: 1.1em !important;
    font-weight: bold !important;
    margin-top: 1.2em;
    margin-bottom: 0.5em;
    color: #666 !important;
}
.privacy-expander-content h3, .disclaimer-expander-content h3 {
    font-size: 0.95em !important;
    font-weight: bold !important;
    margin-top: 1em;
    margin-bottom: 0.4em;
    color: #666 !important;
}
.privacy-expander-content, .disclaimer-expander-content {
    font-size: 0.9em;
    color: #666 !important;
    max-width: 800px;
    margin: auto;
}
.privacy-expander-content p, .disclaimer-expander-content p {
    color: #666 !important;
}
.privacy-expander-content li, .disclaimer-expander-content li {
    color: #666 !important;
}
.privacy-expander-content strong, .disclaimer-expander-content strong {
    color: #666 !important;
}
.privacy-expander-content *, .disclaimer-expander-content * {
    color: #666 !important;
}
/* Make all text in expander gray */
.stExpander [data-testid="stMarkdownContainer"] {
    color: #666 !important;
}
.stExpander [data-testid="stMarkdownContainer"] p,
.stExpander [data-testid="stMarkdownContainer"] li,
.stExpander [data-testid="stMarkdownContainer"] h2,
.stExpander [data-testid="stMarkdownContainer"] h3,
.stExpander [data-testid="stMarkdownContainer"] strong {
    color: #666 !important;
}
</style>
"""
st.markdown(footer_css, unsafe_allow_html=True)

# Privacy Policy and Disclaimer buttons - horizontal, elegant layout
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    btn1, btn2 = st.columns(2)
    with btn1:
        if st.button("Privacy Policy", key="vacation_privacy_btn", help="View Privacy Policy", use_container_width=True):
            st.session_state['show_privacy'] = True
            st.rerun()
    with btn2:
        if st.button("Disclaimer", key="vacation_disclaimer_btn", help="View Disclaimer", use_container_width=True):
            st.session_state['show_disclaimer'] = True
            st.rerun()

# Copyright notice
st.markdown('<div class="footer-copyright">© 2025 Vishnu Balraj</div>', unsafe_allow_html=True)

# Privacy Policy content
PRIVACY_POLICY_TEXT = """
## Privacy Policy

This document outlines the privacy practices for the Vacation Finder & Planner application.

## Data Collection and Storage

This application is designed with your privacy in mind.

*   **API Usage:** This application uses OpenAI API and SerpAPI to generate travel itineraries and search for deals. Your travel preferences, destinations, and dates are sent to these services to provide you with personalized vacation plans.

*   **No Permanent Storage:** The application does not permanently store your travel preferences, itineraries, or personal information on any server. Data is processed in real-time and only maintained in your browser session.

*   **Session State:** The application uses Streamlit's session state functionality to maintain your itinerary and preferences during your active session. This data is cleared when you close your browser tab.

*   **No Analytics or Tracking:** This application does not use tracking services like Google Analytics. We do not collect anonymous usage statistics, browser information, or location data beyond what you explicitly provide.

*   **API Keys:** If you deploy this application yourself, your API keys are stored securely using Streamlit's secrets management and are never exposed to users.

## Your Data

Your travel search queries, preferences, and generated itineraries are processed through third-party APIs (OpenAI and SerpAPI) but are not stored by this application. You are responsible for saving your generated vacation plans to your own device if you wish to keep them.

## Third-Party Services

*   **OpenAI:** Used for generating AI-powered vacation itineraries. Please review OpenAI's privacy policy for information about how they handle data.
*   **SerpAPI:** Used for searching travel deals and information. Please review SerpAPI's privacy policy for information about how they handle data.

---

© 2025 Vishnu Balraj
"""

# Disclaimer content
DISCLAIMER_TEXT = """
**Disclaimer:**

The creators of the Vacation Finder & Planner application make every effort to provide accurate and useful travel information, but we make no guarantees about:

*   The accuracy, completeness, or timeliness of travel deals, prices, hotel recommendations, or restaurant suggestions.
*   The availability of hotels, restaurants, or activities at the recommended locations.
*   The suitability of any recommendations for your specific needs, preferences, or dietary requirements.

**Important Notes:**

*   All prices, availability, and information are subject to change without notice.
*   You should verify all travel information, including hotel availability, restaurant hours, and activity details, directly with the providers before making reservations.
*   The AI-generated itineraries are suggestions only and should be reviewed and customized according to your preferences and needs.
*   The creators are not responsible for any booking issues, travel disruptions, or dissatisfaction with recommended services.
*   Always check current travel advisories, visa requirements, and health guidelines for your destination before traveling.

**Limitation of Liability:**

The creators of this application are not liable for any loss, damage, or inconvenience arising from the use of this application or reliance on any information, recommendations, or links provided herein. Users assume all responsibility for their travel planning and decisions.

---

© 2025 Vishnu Balraj
"""

# Display Privacy Policy if requested
if st.session_state.get('show_privacy', False):
    with st.expander("Privacy Policy", expanded=True):
        privacy_md = PRIVACY_POLICY_TEXT
        privacy_md = re.sub(r'^#+\s*Privacy Policy\s*\n', '', privacy_md, flags=re.IGNORECASE)
        st.markdown('<div class="privacy-main-title">Privacy Policy</div>', unsafe_allow_html=True)
        st.markdown('<div class="privacy-expander-content">', unsafe_allow_html=True)
        st.markdown(privacy_md, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        if st.button("Close", key="close_vacation_privacy"):
            st.session_state['show_privacy'] = False
            st.rerun()

# Display Disclaimer if requested
if st.session_state.get('show_disclaimer', False):
    with st.expander("Disclaimer", expanded=True):
        st.markdown('<div class="disclaimer-main-title">Disclaimer</div>', unsafe_allow_html=True)
        st.markdown('<div class="disclaimer-expander-content">', unsafe_allow_html=True)
        st.markdown(DISCLAIMER_TEXT, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        if st.button("Close", key="close_vacation_disclaimer"):
            st.session_state['show_disclaimer'] = False
            st.rerun()